import os
import io
import tempfile
import boto3
from botocore.exceptions import ClientError
from datetime import datetime
from sqlalchemy import desc
from flask import (
    render_template,
    Blueprint,
    flash,
    redirect,
    url_for,
    request,
    jsonify,
    current_app,
    send_from_directory,
    abort,
    send_file,
    session,
)
from flask_login import login_required, current_user
from flaskr import db, bcrypt
from flaskr.models import Student, Result, SchoolSession, Subject
from flaskr.auths.utils import requires_role, save_and_resize_picture, upload_to_s3
from flaskr.students.forms import (
    StudentRegisterationForm,
    StudentScoreTargetForm,
    StudentResultOptionForm,
    StudentUpdateForm,
    SubjectScoresOptionForm,
)
from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_BREAK

# save_picture

students = Blueprint("students", __name__)

S3_BUCKET = os.environ.get("S3_BUCKET")
S3_REGION = os.environ.get("S3_REGION")
s3 = boto3.client(
    "s3",
    aws_access_key_id=os.environ.get("AWS_ACCESS_KEY_ID"),
    aws_secret_access_key=os.environ.get("AWS_SECRET_ACCESS_KEY"),
    region_name=S3_REGION,
)


def current_session():
    current_session = SchoolSession.query.order_by(
        desc(SchoolSession.start_date)
    ).first()
    current_session = current_session.session if current_session else "2024/2025"
    return current_session


def student_identification(student_class):
    student = Student.query.order_by(desc(Student.date_joined)).first()
    if student:
        recent_student_serial = int(student.sID[11:])
        student_serial = recent_student_serial + 1
    else:
        student_serial = 1
    if student_serial < 10:
        sID = "ISMC2024" + student_class + "00" + str(student_serial)
    elif student_serial < 100:
        sID = "ISMC2024" + student_class + "0" + str(student_serial)
    else:
        sID = "ISMC2024" + student_class + str(student_serial)
    return sID


def grade_calculator(score):
    if score >= 70:
        grade = "A"
    elif score >= 60:
        grade = "B"
    elif score >= 50:
        grade = "C"
    elif score >= 40:
        grade = "D"
    elif score < 40:
        grade = "F"
    else:
        grade = "ungraded"
    return grade


def cummulative_score_calc(results):
    total_score = 0
    for result in results:
        if result.total_score:
            total_score += int(result.total_score)
    return total_score


def distinction_calc(results):
    number_of_distinction = 0
    for result in results:
        if int(result.total_score) >= 70:
            number_of_distinction += 1
    return number_of_distinction


def teacher_comment_func(total_score):
    if total_score >= 70:
        comment = "An excellent perfomance"
    elif total_score >= 60:
        comment = "A good performance"
    elif total_score >= 50:
        comment = "An average performance"
    elif total_score < 50:
        comment = "A poor performance"
    return comment


def interpretation_func(score):
    if score >= 70:
        comment = "EXCELLENT"
    elif score >= 60:
        comment = "VERY GOOD"
    elif score >= 50:
        comment = "CREDIT"
    elif score >= 36:
        comment = "PASS"
    elif score < 36:
        comment = "FAIL"
    return comment


@students.route("/student-dashboard")
@login_required
@requires_role("admin")
def student_dashboard():
    return render_template("student-dashboard.html")


# TODO: generate picture edit and save path
@students.route("/register-student", methods=["GET", "POST"])
@login_required
@requires_role("admin")
def student_register():
    form = StudentRegisterationForm()
    if form.validate_on_submit():
        if form.picture.data:
            try:
                picture_fn, temp_file_path = save_and_resize_picture(form.picture.data)
                s3_url = upload_to_s3(temp_file_path, picture_fn, S3_BUCKET, S3_REGION)
                print(f"Image uploaded successfully! URL: {s3_url}", "success")
            except Exception as e:
                print(f"Error uploading image: {e}", "danger")

            # picture = save_picture(form.picture.data)
            hashed_password = bcrypt.generate_password_hash(
                form.last_name.data.lower()
            ).decode("utf-8")
            sID = student_identification(form.student_class.data)
            student = Student(
                first_name=form.first_name.data,
                middle_name=form.middle_name.data,
                last_name=form.last_name.data,
                dob=form.dob.data,
                sex=form.sex.data,
                email=form.email.data,
                student_class=form.student_class.data,
                department=form.department.data,
                parent_number=form.parent_number.data,
                address=form.address.data,
                picture=picture_fn,
                password=hashed_password,
                sID=sID,
            )
        else:
            hashed_password = bcrypt.generate_password_hash(
                form.last_name.data.lower()
            ).decode("utf-8")
            sID = student_identification(form.student_class.data)
            student = Student(
                first_name=form.first_name.data,
                middle_name=form.middle_name.data,
                last_name=form.last_name.data,
                dob=form.dob.data,
                sex=form.sex.data,
                email=form.email.data,
                student_class=form.student_class.data,
                department=form.department.data,
                parent_number=form.parent_number.data,
                address=form.address.data,
                password=hashed_password,
                sID=sID,
            )
        db.session.add(student)
        db.session.commit()
        flash(
            f"{form.first_name.data} {form.last_name.data} has been successfully registered as a student!",
            "success",
        )
        return redirect(url_for("students.student_register"))
    return render_template("register-student.html", form=form)


@students.route("/get-students/<student_info>")
@login_required
@requires_role("admin")
def get_students(student_info):
    student_info = student_info
    if student_info == "all-students":
        students = Student.query.order_by(
            Student.last_name, Student.first_name, Student.middle_name
        ).all()
        class_range = "ALL"
    elif student_info == "js-1":
        students = (
            Student.query.filter_by(student_class="JS1")
            .order_by(Student.last_name, Student.first_name, Student.middle_name)
            .all()
        )
        class_range = "JSS-1"
    elif student_info == "js-2":
        students = (
            Student.query.filter_by(student_class="JS2")
            .order_by(Student.last_name, Student.first_name, Student.middle_name)
            .all()
        )
        class_range = "JSS-2"
    elif student_info == "js-3":
        students = (
            Student.query.filter_by(student_class="JS3")
            .order_by(Student.last_name, Student.first_name, Student.middle_name)
            .all()
        )
        class_range = "JSS-3"
    elif student_info == "ss-1":
        students = (
            Student.query.filter_by(student_class="SS1")
            .order_by(Student.last_name, Student.first_name, Student.middle_name)
            .all()
        )
        class_range = "SSS-1"
    elif student_info == "ss-2":
        students = (
            Student.query.filter_by(student_class="SS2")
            .order_by(Student.last_name, Student.first_name, Student.middle_name)
            .all()
        )
        class_range = "SSS-2"
    elif student_info == "ss-3":
        students = (
            Student.query.filter_by(student_class="SS3")
            .order_by(Student.last_name, Student.first_name, Student.middle_name)
            .all()
        )
        class_range = "SSS-3"
    else:
        return redirect(url_for("students.student_dashboard"))
    return render_template("students.html", students=students, class_range=class_range)


@students.route("/student-profile/<int:std_id>")
@login_required
@requires_role("student")
def student_profile(std_id):
    student = Student.query.get(std_id)
    year = datetime.now().year
    return render_template("student-profile.html", student=student, year=year)


@students.route("/score-upload-option", methods=["GET", "POST"])
@login_required
@requires_role("teacher")
def score_upload_option():
    form = StudentScoreTargetForm()
    form.subject.choices += Subject.school_subject()
    sessions = SchoolSession.query.order_by(desc(SchoolSession.start_date)).all()
    if sessions:
        for each_session in sessions:
            form.session.choices.append((each_session.session, each_session.session))
    form.session.choices.append(("2024/2025", "2024/2025"))
    if form.validate_on_submit():
        session_val = form.session.data
        term = form.term.data
        subject = form.subject.data
        student_class = form.student_class.data
        department = form.department.data
        if form.session.data == "current_session":
            session_val = current_session()
        if (
            student_class.upper() == "JS1"
            or student_class.upper() == "JS2"
            or student_class.upper() == "JS3"
            or department == "none"
        ):
            department = None
        if department == "all" or department == None:
            students = Student.query.filter_by(student_class=student_class).all()
        else:
            students = Student.query.filter_by(
                student_class=student_class, department=department
            ).all()
        results = Result.query.filter_by(
            subject=subject.lower().replace("-", " "), term=term, session=session_val
        ).all()

        if current_user.role == "teacher":
            return render_template(
                "teacher-score-upload.html",
                teacher=current_user,
                students=students,
                session_val=session_val,
                term=term,
                subject=subject.replace("-", " "),
                student_class=student_class,
                department=department,
                results=results,
            )
        return render_template(
            "student-score-upload.html",
            students=students,
            session_val=session_val,
            term=term,
            subject=subject.replace("-", " "),
            student_class=student_class,
            department=department,
            results=results,
        )
    session_val = session.get("session_val")
    if session_val:
        form.session.data = session_val
    term = session.get("term")
    if term:
        form.term.data = term
    subject = session.get("subject")
    if subject:
        form.subject.data = subject
    student_class = session.get("student_class")
    if student_class:
        form.student_class.data = student_class
    department = session.get("department")
    if department:
        form.department.data = department
    if current_user.role == "teacher":
        return render_template(
            "teacher-upload-option.html", form=form, teacher=current_user
        )
    return render_template("student-upload-option.html", form=form)


@students.route("/score-upload", methods=["POST"])
@login_required
@requires_role("teacher")
def score_upload():
    session_val = request.form.get("session", None)
    term = request.form.get("term")
    subject = request.form.get("subject")
    student_class = request.form.get("student_class")
    department = request.form.get("department", None)
    results = Result.query.filter_by(
        subject=subject.lower().replace("-", " "), term=term, session=session_val
    ).all()
    student_info_exist = False
    if department.lower() == "none" or department.lower() == "all":
        department = None
    if department:
        students = Student.query.filter_by(
            student_class=student_class, department=department
        ).all()
    else:
        students = Student.query.filter_by(student_class=student_class).all()
    for student in students:
        test_score = request.form['test_'+str(student.id)]
        if department == "science":
            practical_score = request.form['practical_'+str(student.id)]
        else:
            practical_score = None
        exam_score = request.form['exam_'+str(student.id)]
        student_info_exist = False
        for result in results:
            if (
                    student.id == result.student_id
                    and term == result.term
                    and session_val == result.session
                    and subject == result.subject
                ):
                print("student exist")
                total_score = int(result.total_score)
                if test_score:
                    if result.test_score:
                        total_score -= int(result.test_score)
                        total_score += int(test_score)
                    result.test_score = int(test_score)
                if exam_score:
                    if result.exam_score:
                        total_score -= int(result.exam_score)
                        total_score += int(exam_score)
                    result.exam_score = int(exam_score)
                if practical_score:
                    if result.practical_score:
                        total_score -= int(result.practical_score)
                        total_score += int(practical_score)
                    result.practical_score = int(practical_score)
                result.total_score = total_score
                student_info_exist = True
                db.session.commit()
                break
        if not student_info_exist:
            print("student not exist")
            total_score = 0
            if test_score:
                total_score += int(test_score)
            if exam_score:
                total_score += int(exam_score)
            if practical_score:
                total_score += int(practical_score)
            result = Result(
                    test_score=int(test_score) if test_score else None,
                    exam_score=int(exam_score) if exam_score else None,
                    practical_score=int(practical_score) if practical_score else None,
                    total_score=total_score,
                    term=term,
                    subject=subject,
                    session=session_val,
                    student_id=int(student.id),
                )
            db.session.add(result)
            db.session.commit()
    flash(
        f"{ student_class } { subject } score have been uploaded!", "success"
    )
    session["session_val"] = session_val
    session["term"] = term
    session["subject"] = subject.lower().replace(" ", "-")
    session["student_class"] = student_class
    session["department"] = department
    return redirect(url_for("students.score_upload_option"))


@students.route("/check-result/<int:std_id>", methods=["GET", "POST"])
@login_required
@requires_role("student")
def result_view(std_id):
    form = StudentResultOptionForm()
    sessions = SchoolSession.query.order_by(desc(SchoolSession.start_date)).all()
    if sessions:
        for each_session in sessions:
            form.session.choices.append((each_session.session, each_session.session))
    form.session.choices.append(("2024/2025", "2024/2025"))
    if request.method == "GET":
        student = Student.query.get(std_id)
        return render_template("student-result-option.html", form=form, student=student)
    elif request.method == "POST":
        if form.validate_on_submit():
            student = Student.query.get(std_id)
            session = form.session.data
            if session.lower() == "current_session":
                session = current_session()
            term = form.term.data
            std_id = int(std_id)
            results = Result.query.filter_by(
                session=session, term=term, student_id=std_id
            ).all()
            return render_template(
                "student-result.html",
                results=results,
                student=student,
                session=session,
                term=term,
            )

    return render_template("student-result-option.html", form=form, student=student)


@students.route("/update-student-profile/<int:std_id>", methods=["GET", "POST"])
@login_required
def update_student_profile(std_id):
    form = StudentUpdateForm()
    student = Student.query.get(std_id)
    if form.validate_on_submit():
        try:
            if form.first_name.data:
                student.first_name = form.first_name.data
            if form.last_name.data:
                student.last_name = form.last_name.data
                hashed_password = bcrypt.generate_password_hash(
                    form.last_name.data.lower()
                ).decode("utf-8")
                student.password = hashed_password
            if form.middle_name.data:
                student.middle_name = form.middle_name.data
            if form.dob.data:
                student.dob = form.dob.data
            if form.sex.data:
                student.sex = form.sex.data
            if form.email.data:
                student.email = form.email.data
            if form.student_class.data:
                student.student_class = form.student_class.data
            if form.department.data:
                student.department = form.department.data
            if form.parent_number.data:
                student.parent_number = form.parent_number.data
            if form.address.data:
                student.address = form.address.data
            if form.picture.data:
                try:
                    image_name = student.picture
                    picture_fn, temp_file_path = save_and_resize_picture(
                        form.picture.data
                    )
                    print(temp_file_path, picture_fn)
                    s3_url = upload_to_s3(
                        temp_file_path, picture_fn, S3_BUCKET, S3_REGION
                    )

                    print(f"Image uploaded successfully! URL: {s3_url}", "success")

                    # deleting former image from bucket
                    if image_name != "default.jpg":
                        try:
                            s3.delete_object(
                                Bucket=S3_BUCKET, Key=f"profile_pics/{image_name}"
                            )
                            print(
                                f"Image {image_name} deleted successfully from {S3_BUCKET}."
                            )
                        except ClientError as e:
                            print(
                                f"Error deleting image {image_name} from {S3_BUCKET}: {e}"
                            )

                except Exception as e:
                    print(f"Error uploading image: {e}", "danger")
                finally:
                    student.picture = picture_fn
            # if form.picture.data:
            #     student.picture = save_picture(form.picture.data)
            db.session.commit()
            flash(
                f"{student.last_name} {student.first_name} profile has been updated successfully!",
                "success",
            )
            return redirect(url_for("students.student_profile", std_id=std_id))
        except:
            flash("There was a problem updating the student profile!", "danger")
            return redirect(url_for("students.student_profile", std_id=std_id))

    form.first_name.data = student.first_name
    form.last_name.data = student.last_name
    form.middle_name.data = student.middle_name
    form.email.data = student.email
    form.student_class.data = student.student_class
    form.department.data = student.department
    form.dob.data = student.dob
    form.address.data = student.address
    form.parent_number.data = student.parent_number
    return render_template("student-update-form.html", form=form)


@students.route("/delete-student/<int:std_id>", methods=["DELETE"])
@login_required
@requires_role("admin")
def delete_student(std_id):
    print("got here")
    student = Student.query.get(std_id)
    image_name = student.picture
    if student:
        try:
            db.session.delete(student)
            db.session.commit()
            # static_dir = os.path.join(current_app.root_path, "static")
            # image_path = os.path.join(static_dir, "profile_pics", f"{image_name}")
            # if os.path.exists(image_path) and image_name != "default.jpg":
            #     os.remove(image_path)

            if image_name != "default.jpg":
                try:
                    s3.delete_object(Bucket=S3_BUCKET, Key=f"profile_pics/{image_name}")
                    print(f"Image {image_name} deleted successfully from {S3_BUCKET}.")
                except ClientError as e:
                    print(f"Error deleting image {image_name} from {S3_BUCKET}: {e}")

            flash(
                f"{student.last_name} {student.first_name} profile has been deleted!",
                "info",
            )
            return render_template("student-dashboard.html")
        except:
            flash(
                f"There was a problem deleting {student.last_name} {student.first_name} profile!",
                "danger",
            )
            return redirect(url_for("students.student_dashboard"))
    flash(
        f"There was a problem deleting {student.last_name} {student.first_name} profile!",
        "danger",
    )
    return redirect(url_for("students.student_dashboard"))


# TODO: Work still on-going here
@students.route("/subject-scores-option", methods=["GET", "POST"])
@login_required
@requires_role("teacher")
def subject_scores_option():
    form = SubjectScoresOptionForm()
    if request.method == "POST":
        student_class = form.student_class.data
        session = form.session.data
        term = form.term.data
        department = form.department.data
        subject = form.subject.data.replace("-", " ")
        print(subject)
        department_exit = False
        if session == "current_session":
            session = current_session()
        if (
            student_class.upper() == "JS1"
            or student_class.upper() == "JS2"
            or student_class.upper() == "JS3"
            or department == "none"
        ):
            department = None
        if not department or department == "all":
            target_results = []
            students = (
                Student.query.filter_by(student_class=student_class)
                .order_by(Student.last_name, Student.first_name)
                .all()
            )
            results = Result.query.filter_by(
                session=session, subject=subject, term=term
            ).all()
            for student in students:
                for result in results:
                    if student.id == result.student_id:
                        target_results.append(result)
        else:
            department_exit = True
            students = (
                Student.query.filter_by(
                    student_class=student_class, department=department
                )
                .order_by(Student.last_name, Student.first_name)
                .all()
            )
            results = Result.query.filter_by(
                session=session, subject=subject, term=term
            ).all()
            target_results = []
            for student in students:
                for result in results:
                    if student.id == result.student_id:
                        target_results.append(result)
        if current_user.role == "teacher":
            temp = "teacher-subject-scores.html"
        elif current_user.role == "admin":
            temp = "subject-scores.html"
        return render_template(
            temp,
            results=target_results,
            subject=subject,
            student_class=student_class,
            session=session,
            term=term,
            department_exit=department_exit,
            department=department,
            teacher=current_user,
        )
    form.subject.choices += Subject.school_subject()
    sessions = SchoolSession.query.order_by(desc(SchoolSession.start_date)).all()
    if sessions:
        for each_session in sessions:
            form.session.choices.append((each_session.session, each_session.session))
    form.session.choices.append(("2024/2025", "2024/2025"))
    if current_user.role == "teacher":
        temp = "teacher-subject-scores-option.html"
    elif current_user.role == "admin":
        temp = "subject-scores-option.html"
    return render_template(temp, form=form, teacher=current_user)


@students.route("/students/search", methods=["GET", "POST"])
@login_required
@requires_role("admin")
def search_student():
    body = []
    try:
        search = request.get_json()["search"]
        student_class = request.get_json()["student_class"]
        print(student_class)
        if student_class.upper() == "ALL":
            students = Student.query.filter(
                Student.last_name.ilike(f"%{search}%")
            ).all()
        elif student_class.upper() == "JSS-1":
            students = (
                Student.query.filter(Student.last_name.ilike(f"%{search}%"))
                .filter_by(student_class="JS1")
                .all()
            )
        elif student_class.upper() == "JSS-2":
            students = (
                Student.query.filter(Student.last_name.ilike(f"%{search}%"))
                .filter_by(student_class="JS2")
                .all()
            )
        elif student_class.upper() == "JSS-3":
            students = (
                Student.query.filter(Student.last_name.ilike(f"%{search}%"))
                .filter_by(student_class="JS3")
                .all()
            )
        elif student_class.upper() == "SSS-1":
            students = (
                Student.query.filter(Student.last_name.ilike(f"%{search}%"))
                .filter_by(student_class="SS1")
                .all()
            )
        elif student_class.upper() == "SSS-2":
            students = (
                Student.query.filter(Student.last_name.ilike(f"%{search}%"))
                .filter_by(student_class="SS2")
                .all()
            )
        elif student_class.upper() == "SSS-3":
            students = (
                Student.query.filter(Student.last_name.ilike(f"%{search}%"))
                .filter_by(student_class="SS3")
                .all()
            )
        count = 1
        for student in students:
            body.append(
                {
                    "id": str(student.id),
                    "serial": str(count),
                    "name": f"{student.last_name} {student.first_name} {student.middle_name}",
                    "sID": student.sID,
                    "studentClass": student.student_class,
                    "department": student.department if student.department else "none",
                }
            )
            count += 1
    except:
        db.session.rollback()
    finally:
        db.session.close()
    return jsonify(body)


# @students.route("/download-result/<int:student_id>", methods=["POST"])
# @login_required
# @requires_role("admin")
# def result_download(student_id):
#     term = request.form.get("term")
#     session = request.form.get("session")
#     student = Student.query.get(student_id)
#     results = Result.query.filter_by(
#         student_id=student_id, term=term, session=session
#     ).all()

#     cummulative_score = cummulative_score_calc(results)
#     expected_total = len(results) * 100
#     percentage = round((cummulative_score / expected_total) * 100, 2)
#     number_of_distinction = distinction_calc(results)
#     teacher_comment = teacher_comment_func(percentage)

#     static_dir = os.path.join(current_app.root_path, "static")
#     template_path = os.path.join(static_dir, "documents/iscore_general_template.docx")
#     template_path_science = os.path.join(
#         static_dir, "documents/iscore_general_template_science.docx"
#     )
#     output_path = os.path.join(static_dir, "documents/student_report_card.docx")

#     if student and results:
#         if student.department.lower() == "science":
#             template = Document(template_path_science)
#         else:
#             template = Document(template_path)
#         template.save(output_path)
#         new_doc = Document(output_path)
#         for paragraph in new_doc.paragraphs:
#             if "{{student_name}}" in paragraph.text:
#                 paragraph.text = paragraph.text.replace(
#                     "{{student_name}}",
#                     f"{student.last_name} {student.first_name} {student.middle_name}",
#                 )
#             elif "{{student_id}}" in paragraph.text:
#                 paragraph.text = paragraph.text.replace(
#                     "{{student_id}}", str(student.sID)
#                 )
#             elif "{{session}}" in paragraph.text:
#                 paragraph.text = paragraph.text.replace("{{session}}", str(session))
#             elif "{{term}}" in paragraph.text:
#                 paragraph.text = paragraph.text.replace("{{term}}", str(term))

#         picture_path = os.path.join(static_dir, f"profile_pics/{student.picture}")
#         table1 = new_doc.tables[0]
#         for row in table1.rows:
#             for cell in row.cells:
#                 if "{{image}}" in cell.text:
#                     cell.text = ""
#                     paragraph = cell.paragraphs[0]
#                     run = paragraph.add_run()
#                     run.add_picture(picture_path, width=Inches(0.8))
#         for table in new_doc.tables:
#             for row in table.rows:
#                 for cell in row.cells:
#                     if "{{student_name}}" in cell.text:
#                         cell.text = cell.text.replace(
#                             "{{student_name}}",
#                             f"{student.last_name} {student.first_name} {student.middle_name}",
#                         )
#                     elif "{{student_class}}" in cell.text:
#                         cell.text = cell.text.replace(
#                             "{{student_class}}",
#                             str(student.student_class.upper()),
#                         )
#                     elif "{{student_id}}" in cell.text:
#                         cell.text = cell.text.replace(
#                             "{{student_id}}",
#                             str(student.sID),
#                         )
#                     elif "{{session}}" in cell.text:
#                         cell.text = cell.text.replace(
#                             "{{session}}",
#                             str(session),
#                         )
#                     elif "{{term}}" in cell.text:
#                         cell.text = cell.text.replace(
#                             "{{term}}",
#                             str(term),
#                         )

#         # Find the table (assuming the first table is where results go)
#         table = new_doc.tables[2]

#         # Append results
#         if student.department.lower() == "science":
#             for result_idx, result in enumerate(results, start=1):
#                 row_cells = table.add_row().cells
#                 row_cells[0].text = str(result_idx)  # S/N
#                 row_cells[1].text = str(result.subject.upper())  # Subject
#                 row_cells[2].text = str(
#                     result.test_score if result.test_score else "-"
#                 )  # Test
#                 row_cells[3].text = str(
#                     result.practical_score if result.practical_score else "-"
#                 )  # Practical
#                 row_cells[4].text = str(
#                     result.exam_score if result.exam_score else "-"
#                 )  # Exam
#                 row_cells[5].text = str(
#                     result.total_score if result.total_score else "-"
#                 )  # Total
#                 row_cells[6].text = grade_calculator(int(result.total_score))  # grade
#                 row_cells[7].text = interpretation_func(
#                     int(result.total_score)
#                 )  # Interpretation
#         else:
#             for result_idx, result in enumerate(results, start=1):
#                 row_cells = table.add_row().cells
#                 row_cells[0].text = str(result_idx)  # S/N
#                 row_cells[1].text = str(result.subject.upper())  # Subject
#                 row_cells[2].text = str(
#                     result.test_score if result.test_score else "-"
#                 )  # Test
#                 row_cells[3].text = str(
#                     result.exam_score if result.exam_score else "-"
#                 )  # Exam
#                 row_cells[4].text = str(
#                     result.total_score if result.total_score else "-"
#                 )  # Total
#                 row_cells[5].text = grade_calculator(int(result.total_score))  # grade
#                 row_cells[6].text = interpretation_func(
#                     int(result.total_score)
#                 )  # Interpretation

#         picture_path_logo = os.path.join(static_dir, "images/iscore_stamp.png")
#         table2 = new_doc.tables[3]
#         for row in table2.rows:
#             for cell in row.cells:
#                 if "{{number_of_subject}}" in cell.text:
#                     cell.text = cell.text.replace(
#                         "{{number_of_subject}}",
#                         str(len(results)),
#                     )
#                 elif "{{cumulative_score}}" in cell.text:
#                     cell.text = cell.text.replace(
#                         "{{cumulative_score}}", str(cummulative_score)
#                     )
#                 elif "{{expected_total}}" in cell.text:
#                     cell.text = cell.text.replace(
#                         "{{expected_total}}", str(expected_total)
#                     )
#                 elif "{{percentage}}" in cell.text:
#                     cell.text = cell.text.replace("{{percentage}}", str(percentage))
#                 elif "{{number_of_distinction}}" in cell.text:
#                     cell.text = cell.text.replace(
#                         "{{number_of_distinction}}", str(number_of_distinction)
#                     )
#                 elif "{{date_printed}}" in cell.text:
#                     parts = cell.text.split("{{date_printed}}")
#                     cell.text = parts[0]
#                     paragraph = cell.paragraphs[0]
#                     new_run = paragraph.add_run(datetime.utcnow().strftime("%d/%m/%Y"))
#                     new_run.font.size = Pt(6)  # Set the font size to 8 points

#         table3 = new_doc.tables[4]
#         for row in table3.rows:
#             for cell in row.cells:
#                 if "{{teachers_comment}}" in cell.text:
#                     cell.text = cell.text.replace(
#                         "{{teachers_comment}}",
#                         teacher_comment,
#                     )

#         new_doc.save(output_path)
#         print(
#             f"Report card for student ID {student.sID} generated and saved to {output_path}"
#         )
#         flash(f"Report card for student ID {student.sID} generated!", "success")
#     else:
#         print(f"Student with ID {student.sID} not found.")
#         flash(f"Student with ID {student.sID} not found.", "info")
#     file_name = "documents/student_report_card.docx"
#     return send_from_directory(static_dir, file_name, as_attachment=True)
# return redirect(url_for("students.student_profile", std_id=str(student_id)))


@students.route("/download-result/<int:student_id>", methods=["POST"])
@login_required
@requires_role("admin")
def result_download(student_id):
    print("Starting result_download function...")
    try:
        # Fetch request data
        term = request.form.get("term")
        session = request.form.get("session")
        print(f"Received term: {term}, session: {session}, student_id: {student_id}")

        # Fetch student and results
        student = Student.query.get(student_id)
        results = Result.query.filter_by(
            student_id=student_id, term=term, session=session
        ).all()
        print(f"Fetched student: {student}, Results count: {len(results)}")

        if not student or not results:
            print("Student or results not found.")
            flash(f"Student with ID {student_id} or results not found.", "info")
            abort(404, description="Student or results not found.")

        # Calculations
        cummulative_score = cummulative_score_calc(results)
        expected_total = len(results) * 100
        percentage = round((cummulative_score / expected_total) * 100, 2)
        number_of_distinction = distinction_calc(results)
        teacher_comment = teacher_comment_func(percentage)
        print(
            f"Calculations completed: Percentage={percentage}, Distinctions={number_of_distinction}"
        )

        # Template paths
        static_dir = os.path.join(current_app.root_path, "static")
        template_path = os.path.join(
            static_dir, "documents/iscore_general_template.docx"
        )
        template_path_science = os.path.join(
            static_dir, "documents/iscore_general_template_science.docx"
        )

        # Choose template based on department
        if student.department.lower() == "science":
            print("Loading science template...")
            template = Document(template_path_science)
        else:
            print("Loading general template...")
            template = Document(template_path)

        # Create a temporary file
        with tempfile.NamedTemporaryFile(delete=False, suffix=".docx") as temp_file:
            output_path = temp_file.name
            print(f"Temporary file created at {output_path}")

        # Save a copy of the template to the temporary file
        template.save(output_path)
        new_doc = Document(output_path)
        print("Replacing placeholders in the document...")

        # Replace placeholders in paragraphs
        for paragraph in new_doc.paragraphs:
            if "{{student_name}}" in paragraph.text:
                paragraph.text = paragraph.text.replace(
                    "{{student_name}}",
                    f"{student.last_name} {student.first_name} {student.middle_name}",
                )
            elif "{{student_id}}" in paragraph.text:
                paragraph.text = paragraph.text.replace(
                    "{{student_id}}", str(student.sID)
                )
            elif "{{session}}" in paragraph.text:
                paragraph.text = paragraph.text.replace("{{session}}", str(session))
            elif "{{term}}" in paragraph.text:
                paragraph.text = paragraph.text.replace("{{term}}", str(term))

        # Handle profile picture replacement
        # picture_path = os.path.join(static_dir, f"profile_pics/{student.picture}")
        # print(f"Profile picture path: {picture_path}")
        image_obj = s3.get_object(
            Bucket=S3_BUCKET, Key=f"profile_pics/{student.picture}"
        )
        print(student.picture)
        image_content = image_obj["Body"].read()
        image_stream = io.BytesIO(image_content)

        table1 = new_doc.tables[0]
        for row in table1.rows:
            for cell in row.cells:
                if "{{image}}" in cell.text:
                    cell.text = ""
                    paragraph = cell.paragraphs[0]
                    run = paragraph.add_run()
                    try:
                        run.add_picture(image_stream, width=Inches(0.8))
                        print("Profile picture added.")
                    except Exception as e:
                        print(f"Error adding profile picture: {e}")

        # Replace placeholders in tables
        for table in new_doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    if "{{student_name}}" in cell.text:
                        cell.text = cell.text.replace(
                            "{{student_name}}",
                            f"{student.last_name} {student.first_name} {student.middle_name}",
                        )
                    elif "{{student_class}}" in cell.text:
                        cell.text = cell.text.replace(
                            "{{student_class}}", str(student.student_class.upper())
                        )
                    elif "{{student_id}}" in cell.text:
                        cell.text = cell.text.replace(
                            "{{student_id}}", str(student.sID)
                        )
                    elif "{{session}}" in cell.text:
                        cell.text = cell.text.replace("{{session}}", str(session))
                    elif "{{term}}" in cell.text:
                        cell.text = cell.text.replace("{{term}}", str(term))

        # Populate result table
        print("Populating result table...")
        table = new_doc.tables[2]
        for result_idx, result in enumerate(results, start=1):
            try:
                row_cells = table.add_row().cells
                row_cells[0].text = str(result_idx)  # S/N
                row_cells[1].text = str(result.subject.upper())  # Subject
                row_cells[2].text = str(result.test_score or "-")  # Test
                if student.department.lower() == "science":
                    row_cells[3].text = str(result.practical_score or "-")  # Practical
                    row_cells[4].text = str(result.exam_score or "-")  # Exam
                    row_cells[5].text = str(result.total_score or "-")  # Total
                    row_cells[6].text = grade_calculator(
                        int(result.total_score)
                    )  # Grade
                    row_cells[7].text = interpretation_func(
                        int(result.total_score)
                    )  # Interpretation
                else:
                    row_cells[3].text = str(result.exam_score or "-")  # Exam
                    row_cells[4].text = str(result.total_score or "-")  # Total
                    row_cells[5].text = grade_calculator(
                        int(result.total_score)
                    )  # Grade
                    row_cells[6].text = interpretation_func(
                        int(result.total_score)
                    )  # Interpretation
            except Exception as e:
                print(f"Error adding result row {result_idx}: {e}")

        table2 = new_doc.tables[3]
        for row in table2.rows:
            for cell in row.cells:
                if "{{number_of_subject}}" in cell.text:
                    cell.text = cell.text.replace(
                        "{{number_of_subject}}",
                        str(len(results)),
                    )
                elif "{{cumulative_score}}" in cell.text:
                    cell.text = cell.text.replace(
                        "{{cumulative_score}}", str(cummulative_score)
                    )
                elif "{{expected_total}}" in cell.text:
                    cell.text = cell.text.replace(
                        "{{expected_total}}", str(expected_total)
                    )
                elif "{{percentage}}" in cell.text:
                    cell.text = cell.text.replace("{{percentage}}", str(percentage))
                elif "{{number_of_distinction}}" in cell.text:
                    cell.text = cell.text.replace(
                        "{{number_of_distinction}}", str(number_of_distinction)
                    )
                elif "{{date_printed}}" in cell.text:
                    parts = cell.text.split("{{date_printed}}")
                    cell.text = parts[0]
                    paragraph = cell.paragraphs[0]
                    new_run = paragraph.add_run(datetime.utcnow().strftime("%d/%m/%Y"))
                    new_run.font.size = Pt(6)  # Set the font size to 8 points

        table3 = new_doc.tables[4]
        for row in table3.rows:
            for cell in row.cells:
                if "{{teachers_comment}}" in cell.text:
                    cell.text = cell.text.replace(
                        "{{teachers_comment}}",
                        teacher_comment,
                    )

        # Finalize document
        output_stream = io.BytesIO()
        new_doc.save(output_stream)
        output_stream.seek(0)

        # new_doc.save(output_path)
        # print(f"Final document saved to {output_path}")

        # Serve the file
        return send_file(
            output_stream,
            as_attachment=True,
            download_name=f"{student.last_name}_{student.first_name}_{student.student_class}_report_card.docx",
        )
    except Exception as e:
        print(f"Error in result_download: {e}")
        abort(500, description="An error occurred while processing the request.")
    finally:
        # Clean up temporary file
        if "output_path" in locals() and os.path.exists(output_path):
            os.remove(output_path)
            print(f"Temporary file {output_path} deleted.")


# TODO: A lot to be done here
@students.route("/download-results", methods=["POST"])
@login_required
@requires_role("admin")
def results_download():
    term = "1st"
    session = "2024/2025"
    student_class = request.form.get("class_range")
    if (
        student_class == None
        or student_class == "ALL"
        or student_class == "all"
        or student_class == "All"
        or student_class == "none"
    ):
        students = Student.query.all()
    else:
        if student_class == "JSS-1":
            student_class = "JS1"
        elif student_class == "JSS-2":
            student_class = "JS2"
        elif student_class == "JSS-3":
            student_class = "JS3"
        elif student_class == "SSS-1":
            student_class = "SS1"
        elif student_class == "SSS-2":
            student_class = "SS2"
        elif student_class == "SSS-3":
            student_class = "SS3"
        students = Student.query.filter_by(student_class=student_class).all()

    # term = request.form.get("term")
    # session = request.form.get("session")
    # class_name = request.form.get("class_range")  # Optional: to filter by class

    # Query all students (optionally filter by class)
    # if class_name:
    #     students = Student.query.filter_by(student_class=class_name).order_by(Student.last_name).all()
    # else:
    #     students = Student.query.order_by(Student.last_name).all()

    static_dir = os.path.join(current_app.root_path, "static")
    template_path = os.path.join(static_dir, "documents/iscore_general_template.docx")
    template_path_science = os.path.join(
        static_dir, "documents/iscore_general_template_science.docx"
    )
    output_path = os.path.join(static_dir, "documents/bulk_report_cards.docx")

    # Create a new document for all reports

    # Query all students (optionally filter by class)

    # Create a new document for all reports
    # term = request.form.get("term")
    # session = request.form.get("session")
    # class_name = request.form.get("class_name")  # Optional: to filter by class

    # Query all students (optionally filter by class)
    # if class_name:
    #     students = Student.query.filter_by(student_class=class_name).order_by(Student.last_name).all()
    # else:
    #     students = Student.query.order_by(Student.last_name).all()

    # Create a new document for all reports
    bulk_doc = Document()

    for idx, student in enumerate(students):
        results = Result.query.filter_by(
            student_id=student.id, term=term, session=session
        ).all()

        if not results:
            continue  # Skip students with no results

        # Calculate student metrics
        cummulative_score = cummulative_score_calc(results)
        expected_total = len(results) * 100
        percentage = round((cummulative_score / expected_total) * 100, 2)
        number_of_distinction = distinction_calc(results)
        teacher_comment = teacher_comment_func(percentage)

        # If not first student, add page break
        if idx > 0:
            bulk_doc.add_paragraph().add_run().add_break(WD_BREAK.PAGE)

        # Load appropriate template based on department
        if student.department.lower() == "science":
            template = Document(template_path_science)
        else:
            template = Document(template_path)

        # Copy template content to bulk document
        for paragraph in template.paragraphs:
            new_paragraph = bulk_doc.add_paragraph()
            # Copy the paragraph text and formatting
            for run in paragraph.runs:
                new_run = new_paragraph.add_run(run.text)
                # Copy run's formatting
                new_run.bold = run.bold
                new_run.italic = run.italic
                new_run.underline = run.underline
                if run.font.size:
                    new_run.font.size = run.font.size

        # Copy tables
        for table in template.tables:
            new_table = bulk_doc.add_table(rows=0, cols=len(table.columns))
            new_table.style = table.style

            for row in table.rows:
                new_row = new_table.add_row()
                for idx, cell in enumerate(row.cells):
                    new_cell = new_row.cells[idx]
                    new_cell.text = cell.text
                    # Copy cell formatting if needed
                    if cell.paragraphs:
                        new_cell._tc = cell._tc

        # Get the section of tables we just added
        start_table_index = len(bulk_doc.tables) - len(template.tables)

        # Replace placeholders in paragraphs
        for paragraph in bulk_doc.paragraphs[-len(template.paragraphs) :]:
            if "{{student_name}}" in paragraph.text:
                paragraph.text = paragraph.text.replace(
                    "{{student_name}}",
                    f"{student.last_name} {student.first_name} {student.middle_name}",
                )
            elif "{{student_id}}" in paragraph.text:
                paragraph.text = paragraph.text.replace(
                    "{{student_id}}", str(student.sID)
                )
            elif "{{session}}" in paragraph.text:
                paragraph.text = paragraph.text.replace("{{session}}", str(session))
            elif "{{term}}" in paragraph.text:
                paragraph.text = paragraph.text.replace("{{term}}", str(term))

        # Handle student image
        picture_path = os.path.join(static_dir, f"profile_pics/{student.picture}")
        table1 = bulk_doc.tables[start_table_index]
        for row in table1.rows:
            for cell in row.cells:
                if "{{image}}" in cell.text:
                    cell.text = ""
                    paragraph = cell.paragraphs[0]
                    run = paragraph.add_run()
                    run.add_picture(picture_path, width=Inches(0.8))

        # Replace placeholders in all tables
        for table in bulk_doc.tables[start_table_index:]:
            for row in table.rows:
                for cell in row.cells:
                    if "{{student_name}}" in cell.text:
                        cell.text = cell.text.replace(
                            "{{student_name}}",
                            f"{student.last_name} {student.first_name} {student.middle_name}",
                        )
                    elif "{{student_class}}" in cell.text:
                        cell.text = cell.text.replace(
                            "{{student_class}}",
                            str(student.student_class.upper()),
                        )
                    elif "{{student_id}}" in cell.text:
                        cell.text = cell.text.replace(
                            "{{student_id}}",
                            str(student.sID),
                        )
                    elif "{{session}}" in cell.text:
                        cell.text = cell.text.replace(
                            "{{session}}",
                            str(session),
                        )
                    elif "{{term}}" in cell.text:
                        cell.text = cell.text.replace(
                            "{{term}}",
                            str(term),
                        )

        # Handle results table
        results_table = bulk_doc.tables[
            start_table_index + 2
        ]  # Assuming it's the third table

        if student.department.lower() == "science":
            for result_idx, result in enumerate(results, start=1):
                row_cells = results_table.add_row().cells
                row_cells[0].text = str(result_idx)
                row_cells[1].text = str(result.subject.upper())
                row_cells[2].text = str(result.test_score if result.test_score else "-")
                row_cells[3].text = str(
                    result.practical_score if result.practical_score else "-"
                )
                row_cells[4].text = str(result.exam_score if result.exam_score else "-")
                row_cells[5].text = str(
                    result.total_score if result.total_score else "-"
                )
                row_cells[6].text = grade_calculator(int(result.total_score))
                row_cells[7].text = interpretation_func(int(result.total_score))
        else:
            for result_idx, result in enumerate(results, start=1):
                row_cells = results_table.add_row().cells
                row_cells[0].text = str(result_idx)
                row_cells[1].text = str(result.subject.upper())
                row_cells[2].text = str(result.test_score if result.test_score else "-")
                row_cells[3].text = str(result.exam_score if result.exam_score else "-")
                row_cells[4].text = str(
                    result.total_score if result.total_score else "-"
                )
                row_cells[5].text = grade_calculator(int(result.total_score))
                row_cells[6].text = interpretation_func(int(result.total_score))

        # Handle summary tables
        summary_table = bulk_doc.tables[start_table_index + 3]
        for row in summary_table.rows:
            for cell in row.cells:
                if "{{number_of_subject}}" in cell.text:
                    cell.text = cell.text.replace(
                        "{{number_of_subject}}",
                        str(len(results)),
                    )
                elif "{{cumulative_score}}" in cell.text:
                    cell.text = cell.text.replace(
                        "{{cumulative_score}}", str(cummulative_score)
                    )
                elif "{{expected_total}}" in cell.text:
                    cell.text = cell.text.replace(
                        "{{expected_total}}", str(expected_total)
                    )
                elif "{{percentage}}" in cell.text:
                    cell.text = cell.text.replace("{{percentage}}", str(percentage))
                elif "{{number_of_distinction}}" in cell.text:
                    cell.text = cell.text.replace(
                        "{{number_of_distinction}}", str(number_of_distinction)
                    )
                elif "{{date_printed}}" in cell.text:
                    parts = cell.text.split("{{date_printed}}")
                    cell.text = parts[0]
                    paragraph = cell.paragraphs[0]
                    new_run = paragraph.add_run(datetime.utcnow().strftime("%d/%m/%Y"))
                    new_run.font.size = Pt(6)

        # Handle teacher's comment
        comment_table = bulk_doc.tables[start_table_index + 4]
        for row in comment_table.rows:
            for cell in row.cells:
                if "{{teachers_comment}}" in cell.text:
                    cell.text = cell.text.replace(
                        "{{teachers_comment}}",
                        teacher_comment,
                    )

    # Save the final document
    bulk_doc.save(output_path)

    # Generate a meaningful filename
    timestamp = datetime.utcnow().strftime("%Y%m%d_%H%M%S")
    download_filename = f"bulk_report_cards_{session}_{term}_{timestamp}.docx"

    return send_from_directory(
        static_dir,
        "documents/bulk_report_cards.docx",
        as_attachment=True,
        download_name=download_filename,
    )
